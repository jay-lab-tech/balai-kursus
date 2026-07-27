import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

src = Path('docs/LAPORAN KEGIATAN PKL RPL (UPI) - RAPI_FIXED.docx')
doc = Document(str(src))

# Find BAB VI paragraph to insert before it
bab6_el = None
for para in doc.paragraphs:
    if para.text.strip() == 'BAB VI' and 'Heading' in (para.style.name or ''):
        bab6_el = para._element
        break

if bab6_el is None:
    print("[ERROR] BAB VI not found")
    sys.exit(1)

# Check if page break exists before BAB VI; find the page-break paragraph
# We insert new content before the page break so BAB VI stays on its own page.
prev = bab6_el.getprevious()
insert_ref = bab6_el
# If previous element is a page break paragraph, insert before it
if prev is not None and prev.findall('.//' + qn('w:br')):
    insert_ref = prev

def make_para_before(ref_el, text, style_name, doc):
    new_p = OxmlElement('w:p')
    ref_el.addprevious(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            p.text = text
            try:
                p.style = style_name
            except:
                pass
            return p
    return None

make_para_before(insert_ref, '5.3 Pengujian Sistem', 'Heading 2', doc)
make_para_before(insert_ref,
    'Pengujian sistem dilakukan untuk memastikan seluruh fitur berjalan sesuai kebutuhan fungsional. '
    'Metode pengujian yang digunakan adalah Black Box Testing, yaitu pengujian yang berfokus pada '
    'masukan dan keluaran sistem tanpa memeriksa struktur kode internal. Hasil pengujian utama '
    'adalah sebagai berikut:',
    'Normal', doc)

# Build a test result table
tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
hdr[0].text = 'No'
hdr[1].text = 'Skenario Pengujian'
hdr[2].text = 'Hasil yang Diharapkan'
hdr[3].text = 'Status'

test_cases = [
    ('1', 'Login dengan akun valid', 'Pengguna masuk ke dashboard sesuai perannya', 'Berhasil'),
    ('2', 'Login dengan akun tidak valid', 'Sistem menolak dan menampilkan pesan error', 'Berhasil'),
    ('3', 'Registrasi peserta baru', 'Akun peserta baru berhasil dibuat', 'Berhasil'),
    ('4', 'CRUD data master (program, level, kursus)', 'Data tersimpan, terubah, dan terhapus dengan benar', 'Berhasil'),
    ('5', 'Penjadwalan dengan jadwal bentrok', 'Sistem menolak jadwal yang bentrok', 'Berhasil'),
    ('6', 'Pencatatan absensi oleh instruktur', 'Status kehadiran peserta tersimpan', 'Berhasil'),
    ('7', 'Input nilai peserta', 'Nilai tersimpan dan terhitung otomatis', 'Berhasil'),
    ('8', 'Pembayaran online via Midtrans', 'Transaksi diproses dan status terupdate', 'Berhasil'),
    ('9', 'Penerbitan sertifikat', 'Sertifikat PDF dengan QR code ter-generate', 'Berhasil'),
    ('10', 'Verifikasi sertifikat via QR code', 'Data sertifikat tampil valid', 'Berhasil'),
]
for row in test_cases:
    cells = tbl.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# Move the table element (currently at end of body) to before insert_ref
tbl_el = tbl._element
tbl_el.getparent().remove(tbl_el)
insert_ref.addprevious(tbl_el)

# caption below table
cap = make_para_before(insert_ref, 'Table 5.1 Hasil Pengujian Black Box', 'Heading 4', doc)

make_para_before(insert_ref,
    'Berdasarkan hasil pengujian di atas, seluruh fitur utama sistem berfungsi sesuai dengan '
    'kebutuhan yang telah ditentukan. Tidak ditemukan kesalahan kritis yang menghambat operasional sistem.',
    'Normal', doc)

make_para_before(insert_ref, '5.4 Deployment', 'Heading 2', doc)
make_para_before(insert_ref,
    'Setelah pengembangan dan pengujian selesai, aplikasi Balai Kursus di-deploy agar dapat diakses '
    'secara online. Proses deployment meliputi langkah-langkah berikut:\n'
    '1. Menyiapkan server dengan spesifikasi yang mendukung PHP 8.1, MySQL, dan web server.\n'
    '2. Meng-clone repository dari GitHub ke server produksi.\n'
    '3. Menjalankan composer install dan npm install untuk memasang dependensi.\n'
    '4. Melakukan konfigurasi file .env (koneksi database, kredensial Midtrans, konfigurasi CAS).\n'
    '5. Menjalankan migrasi database (php artisan migrate) dan build aset frontend (npm run build).\n'
    '6. Mengatur permission folder storage dan bootstrap/cache.\n'
    '7. Melakukan optimasi dengan php artisan config:cache dan route:cache.\n'
    'Setelah proses tersebut, aplikasi berhasil berjalan secara online dan dapat diakses oleh '
    'pengguna melalui browser.',
    'Normal', doc)

doc.save(str(src))
print("[DONE] Added 5.3 Pengujian and 5.4 Deployment")
