from docx import Document
from pathlib import Path

p = Path('docs/LAPORAN KEGIATAN PKL RPL (UPI) - RAPI.docx')
doc = Document(str(p))

# Print all headings with their paragraph index AND first few lines of content after them
print('=== FULL HEADING + CONTENT PREVIEW ===')
for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    style = para.style.name if para.style else '?'
    if style.startswith('Heading') or style == 'heading' or style == 'Title':
        print(f'\n[{i}] <{style}> {txt}')
        # Show next 3 non-empty paragraphs after this heading
        count = 0
        for j in range(i+1, min(i+10, len(doc.paragraphs))):
            nxt = doc.paragraphs[j].text.strip()
            nxt_style = doc.paragraphs[j].style.name
            if nxt_style.startswith('Heading') or nxt_style == 'heading':
                break
            if nxt and count < 3:
                print(f'  -> [{j}] {nxt[:100]}')
                count += 1
