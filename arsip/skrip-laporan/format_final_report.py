from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

INPUT = Path('docs/LAPORAN KEGIATAN PKL RPL (UPI) - FINAL.docx')
OUTPUT = Path('docs/LAPORAN KEGIATAN PKL RPL (UPI) - FINAL.docx')

doc = Document(str(INPUT))

def set_font(run, name='Times New Roman', size=12, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def clear_para(p):
    for child in list(p._element):
        if child.tag != qn('w:pPr'):
            p._element.remove(child)

def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), instruction)
    run._r.addnext(fld)

def insert_before(ref, text='', style=None, page_break=False, align=None):
    el = OxmlElement('w:p')
    ref._element.addprevious(el)
    p = next(p for p in doc.paragraphs if p._element is el)
    if style:
        p.style = style
    if page_break:
        p.paragraph_format.page_break_before = True
    if text:
        p.add_run(text)
    if align is not None:
        p.alignment = align
    return p

# Required page geometry: A4, top 3 cm, bottom 2.5 cm, left 4 cm, right 2.5 cm.
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(4)
    section.right_margin = Cm(2.5)

# Set the main styles explicitly instead of relying on Word defaults.
for style_name in ['Normal', 'Normal (Web)', 'List Paragraph']:
    try:
        st = doc.styles[style_name]
        st.font.name = 'Times New Roman'
        st._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Times New Roman')
        st._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        st.font.size = Pt(12)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_after = Pt(0)
    except KeyError:
        pass

for style_name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
    try:
        st = doc.styles[style_name]
        st.font.name = 'Times New Roman'
        st._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Times New Roman')
        st._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        st.font.size = Pt(14)
        st.font.bold = True
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)
    except KeyError:
        pass

# Correct known statements so the finished report matches the actual project.
replacements = {
    'Pengecualian Transaksi Keuangan: Sistem tidak mencakup modul atau proses pembayaran secara daring (online payment gateway). Proses keuangan dilakukan di luar sistem, sehingga aplikasi hanya fokus pada administrasi pendaftaran dan verifikasi.':
    'Cakupan Transaksi Keuangan: Sistem mencakup pembayaran daring melalui Midtrans dan pemutakhiran status pembayaran pendaftaran. Rekonsiliasi akuntansi dan pelaporan keuangan di luar ruang lingkup project.',
    'Pembayaran online terintegrasi dengan Midtrans.':
    'Pembayaran daring terintegrasi dengan Midtrans, termasuk pembuatan transaksi dan pemutakhiran status melalui callback/webhook.',
    'serta memantau analitik penerbitan sertifikat.':
    'serta mengelola status sertifikat. Fitur analitik khusus belum tersedia pada source project yang direview.',
    'Cetak Bukti Pendaftaran: Menghasilkan berkas PDF slip atau kartu peserta setelah pendaftaran berhasil.':
    'Pembuatan PDF Sertifikat: Menghasilkan dokumen PDF sertifikat berdasarkan data dan template sertifikat.',
    'Fungsi: Menghasilkan kode QR otomatis pada bukti pendaftaran sebagai media validasi data peserta saat presensi kursus.':
    'Fungsi: Mendukung kode/QR verifikasi pada sertifikat agar pihak luar dapat memeriksa keaslian sertifikat.',
    'serta mengimpor data massal secara instan.':
    'serta mengekspor data peserta dan nilai ke format Excel.',
    'Fungsi: Melakukan proses pengiriman data formulir pendaftaran dan pengambilan informasi secara real-time (AJAX) tanpa perlu memuat ulang halaman.':
    'Fungsi: Mendukung kebutuhan HTTP asynchronous pada frontend jika digunakan oleh komponen antarmuka.',
}
for p in doc.paragraphs:
    for old, new in replacements.items():
        if old in p.text:
            p.text = p.text.replace(old, new)

# Add TOC/list fields after the existing DAFTAR ISI heading.
toc_heading = next((p for p in doc.paragraphs if p.text.strip().upper() == 'DAFTAR ISI'), None)
if toc_heading is not None and 'TOC \\o' not in doc.element.xml:
    toc = insert_before(next((p for p in doc.paragraphs if p._element is toc_heading._element.getnext()), toc_heading), style='Normal')
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    toc.paragraph_format.page_break_after = True

# Add explicit list headings if they are absent. Word updates these fields when opened.
all_text = '\n'.join(p.text.upper() for p in doc.paragraphs)
anchor = next((p for p in doc.paragraphs if p.text.strip().upper() == 'BAB I'), None)
if anchor is not None and 'DAFTAR TABEL' not in all_text:
    p = insert_before(anchor, 'DAFTAR TABEL', 'Heading 1', page_break=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    f = insert_before(anchor, style='Normal')
    add_field(f, 'TOC \\t "Table 4.1 Tabel users,Table 4.2 Tabel pesertas,Table 4.3 Tabel instrukturs" \\h \\z \\u')
    p = insert_before(anchor, 'DAFTAR GAMBAR', 'Heading 1', page_break=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    f = insert_before(anchor, style='Normal')
    add_field(f, 'TOC \\a "Figure" \\h \\z \\u')

# Add the required bibliography and appendix sections before BAB VI.
chapter6 = next((p for p in doc.paragraphs if ' '.join(p.text.split()) == 'BAB VI'), None)
all_text = '\n'.join(p.text.upper() for p in doc.paragraphs)
if chapter6 is not None and 'DAFTAR PUSTAKA' not in all_text:
    insert_before(chapter6, 'Daftar Pustaka', 'Heading 1', page_break=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    insert_before(chapter6,
        'Laravel. (2023). Laravel 10.x Documentation. https://laravel.com/docs/10.x\n'
        'Midtrans. (2024). Midtrans Documentation. https://docs.midtrans.com/\n'
        'Tailwind CSS. (2024). Tailwind CSS Documentation. https://tailwindcss.com/docs\n'
        'Vite. (2024). Vite Documentation. https://vite.dev/guide/\n'
        'Universitas Pendidikan Indonesia. (2024). Informasi Universitas Pendidikan Indonesia. https://www.upi.edu/',
        'Normal')
    insert_before(chapter6, 'Lampiran', 'Heading 1', page_break=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    insert_before(chapter6, 'Lampiran A. Daftar Dokumentasi Tampilan', 'Heading 2')
    insert_before(chapter6,
        'Lampiran ini memuat dokumentasi tampilan fitur utama aplikasi yang digunakan sebagai bukti implementasi, '
        'meliputi papan informasi, autentikasi, dashboard Admin/Instruktur/Peserta, master data, jadwal, risalah, '
        'absensi, nilai, pembayaran, dan sertifikat. Screenshot terkait telah ditempatkan pada Bab V.',
        'Normal')
    insert_before(chapter6, 'Lampiran B. Ringkasan Skenario Pengujian', 'Heading 2')
    insert_before(chapter6,
        'Skenario pengujian mencakup autentikasi, pembatasan role, pendaftaran dan placement, pembayaran Midtrans, '
        'validasi konflik jadwal, pengelolaan pembelajaran, ekspor nilai/peserta, serta penerbitan dan verifikasi sertifikat. '
        'Rincian ringkas hasil pengujian tersedia pada Subbab 5.3.',
        'Normal')

# Apply body/heading formatting consistently and add page starts for chapters.
chapter_names = {'BAB I', 'BAB II', 'BAB III', 'BAB IV', 'BAB V', 'BAB VI'}
for p in doc.paragraphs:
    text = ' '.join(p.text.split())
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if text in chapter_names:
        p.style = 'Heading 1'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.page_break_before = True
        for r in p.runs:
            set_font(r, size=14, bold=True)
        continue
    if text.startswith('Figure ') or text.startswith('Table '):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_font(r, size=12, bold=True)
        continue
    is_heading = p.style.name.lower().startswith('heading') or text in {
        'Profil Perusahaan', 'Kegiatan Praktik Kerja Lapangan (PKL)',
        'Implementasi dan Pengujian', 'Penutup'
    }
    if is_heading:
        for r in p.runs:
            set_font(r, size=14, bold=True)
    else:
        for r in p.runs:
            set_font(r, size=12)

# Make the cover title and institutional lines explicit 14 pt bold.
for p in doc.paragraphs[:5]:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        set_font(r, size=14, bold=True)

# Make Word refresh generated fields when opened in Microsoft Word.
settings = doc.settings.element
update = settings.find(qn('w:updateFields'))
if update is None:
    update = OxmlElement('w:updateFields')
    settings.append(update)
update.set(qn('w:val'), 'true')

doc.save(str(OUTPUT))
print(f'Saved {OUTPUT}')
