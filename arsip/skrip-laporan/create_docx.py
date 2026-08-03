import re
from docx import Document
from pathlib import Path

source = Path('docs/BAB_IV_PEMBAHASAN.txt')
target = Path('docs/BAB_IV_PEMBAHASAN_new.docx')
text = source.read_text(encoding='utf-8')
lines = text.splitlines()

doc = Document()
current_table = None

def close_table():
    global current_table
    current_table = None

for line in lines:
    stripped = line.strip()
    if stripped == '':
        close_table()
        doc.add_paragraph('')
        continue

    if stripped.startswith('BAB '):
        close_table()
        p = doc.add_paragraph(stripped)
        p.style = 'Heading 1'
        continue

    if re.match(r'^[0-9]+\.[0-9]+\.[0-9]+ Table ', stripped):
        close_table()
        p = doc.add_paragraph(stripped)
        p.style = 'Heading 3'
        continue

    if re.match(r'^[0-9]+\.[0-9]+ Table ', stripped):
        close_table()
        p = doc.add_paragraph(stripped)
        p.style = 'Heading 2'
        continue

    if stripped.startswith('- '):
        if current_table is None:
            current_table = doc.add_table(rows=1, cols=2)
            hdr_cells = current_table.rows[0].cells
            hdr_cells[0].text = 'Field'
            hdr_cells[1].text = 'Keterangan'
        item = stripped[2:].strip()
        if ':' in item:
            field, desc = item.split(':', 1)
            row_cells = current_table.add_row().cells
            row_cells[0].text = field.strip()
            row_cells[1].text = desc.strip()
        else:
            row_cells = current_table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = ''
        continue

    close_table()
    doc.add_paragraph(stripped)

doc.save(target)
