import sys
import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
from pathlib import Path
from copy import deepcopy

sys.stdout.reconfigure(encoding='utf-8')

source = Path('docs/LAPORAN KEGIATAN PKL RPL (UPI) - RAPI.docx')
backup = Path('docs/LAPORAN KEGIATAN PKL RPL (UPI) - RAPI_BACKUP.docx')
output = Path('docs/LAPORAN KEGIATAN PKL RPL (UPI) - FINAL.docx')

if not backup.exists():
    shutil.copy(source, backup)
    print("[OK] Backup created")

doc = Document(str(source))
paras = doc.paragraphs
body = doc.element.body
print(f"[OK] Loaded: {len(paras)} paragraphs")

# -------------------------------------------------------------------
# Helper: add a new paragraph BEFORE a given paragraph element
# -------------------------------------------------------------------
def insert_para_before(ref_para, text, style_name):
    """Insert a new paragraph before ref_para in the body."""
    new_para = OxmlElement('w:p')
    new_run = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_run.append(new_t)
    new_para.append(new_run)
    ref_para._element.addprevious(new_para)
    # Apply style via python-docx paragraph wrapper
    # We need to get the paragraph object; find it in doc
    for p in doc.paragraphs:
        if p._element is new_para:
            try:
                p.style = style_name
            except Exception:
                pass
            return p
    return None

# -------------------------------------------------------------------
# Helper: change heading text and style of an existing paragraph
# -------------------------------------------------------------------
def set_heading(para, text, style_name):
    para.clear()
    para.text = text
    try:
        para.style = style_name
    except Exception:
        pass

# -------------------------------------------------------------------
# Key indices (from diagnostic):
#   99 = PROFIL PERUSAHAAN
#   185 = Landasan teori
#   322 = BAB IV heading (style 'heading')
#   323 = 'Pembahasan' subheading (style 'heading')
#   324 = 'Tabel Database yang Digunakan' (Heading 2)
#   539 = 'Implementasi Sistem' (Heading 2)
#   568 = 'Pembahasan dan Hasil' (Heading 2)
#   752 = last paragraph (Figure 4.38)
# -------------------------------------------------------------------
bab4_para    = paras[322]   # BAB IV
pembahasan_sub = paras[323] # "Pembahasan"
tabel_db_para  = paras[324] # Tabel Database yang Digunakan
implementasi_para = paras[539]
pembahasan_hasil_para = paras[568]
profil_para = paras[99]
landasan_para = paras[185]

print("[INFO] Step 1: Rename BAB IV -> BAB IV (Analisis & Rancangan Proyek)")
# Change BAB IV heading text
for run in bab4_para.runs:
    run.text = ''
bab4_para.clear()
bab4_para.text = 'BAB IV'
try:
    bab4_para.style = 'Heading 1'
except:
    pass

# Change "Pembahasan" subheading -> subtitle for BAB IV
for run in pembahasan_sub.runs:
    run.text = ''
pembahasan_sub.clear()
pembahasan_sub.text = 'Analisis dan Rancangan Proyek'
try:
    pembahasan_sub.style = 'Heading 1'
except:
    pass

print("[INFO] Step 2: Add sub-bab for BAB IV before Tabel Database")
# Insert: 4.1 Identifikasi Masalah, 4.2 Kebutuhan Sistem, 4.3 ERD, 4.4 Perancangan Interface
# before Tabel Database (which becomes 4.5 Struktur Database)

# We'll insert BAB IV sub-bab sections before tabel_db_para
ref = tabel_db_para._element

def make_para_before(ref_el, text, style_name, doc):
    """Create and insert a paragraph before ref_el, return it."""
    new_p = OxmlElement('w:p')
    ref_el.addprevious(new_p)
    # find wrapper
    for p in doc.paragraphs:
        if p._element is new_p:
            p.text = text
            try:
                p.style = style_name
            except:
                pass
            return p
    return None

make_para_before(ref, '4.1 Identifikasi Masalah', 'Heading 2', doc)
make_para_before(ref,
    'Permasalahan utama yang ditemukan selama PKL adalah sistem pengelolaan kursus di DSTI UPI '
    'masih dilakukan secara manual atau menggunakan sistem lama yang belum terintegrasi. Proses '
    'pendaftaran peserta, penjadwalan kursus, pencatatan absensi, penilaian, hingga penerbitan '
    'sertifikat memerlukan waktu yang lama dan rawan kesalahan. Berdasarkan observasi dan diskusi '
    'dengan pembimbing lapangan, identifikasi masalah adalah sebagai berikut:\n'
    '1. Pendaftaran kursus masih dilakukan secara offline/manual sehingga tidak efisien.\n'
    '2. Tidak ada sistem terintegrasi untuk mengelola peserta, instruktur, jadwal, dan pembayaran.\n'
    '3. Penerbitan sertifikat dilakukan secara manual tanpa mekanisme verifikasi digital.\n'
    '4. Tidak ada dashboard terpusat untuk memantau aktivitas kursus secara real-time.',
    'Normal', doc)
make_para_before(ref, '4.2 Kebutuhan Sistem', 'Heading 2', doc)
make_para_before(ref,
    'Berdasarkan identifikasi masalah, kebutuhan fungsional sistem yang dirancang adalah sebagai berikut:\n'
    '1. Sistem pendaftaran kursus online yang dapat diakses peserta dari mana saja.\n'
    '2. Manajemen data master: program, level, kursus, lokasi, kelas, hari, instruktur, peserta.\n'
    '3. Penjadwalan kursus dengan validasi benturan jadwal.\n'
    '4. Pencatatan risalah (berita acara) dan absensi setiap pertemuan oleh instruktur.\n'
    '5. Sistem penilaian peserta dengan komponen nilai yang dapat dikonfigurasi.\n'
    '6. Pembayaran online terintegrasi dengan Midtrans.\n'
    '7. Penerbitan sertifikat digital dengan QR code untuk verifikasi publik.\n'
    '8. Dashboard berbasis peran (Admin, Instruktur, Peserta).',
    'Normal', doc)
make_para_before(ref, '4.3 Entity Relationship Diagram (ERD)', 'Heading 2', doc)
make_para_before(ref,
    'ERD berikut menggambarkan relasi antar-entitas utama dalam sistem Balai Kursus. '
    'Entitas utama meliputi users, pesertas, instrukturs, programs, levels, kursuses, jadwals, '
    'pendaftarans, payments, risalahs, absensis, scores, dan certificates. '
    'Relasi antar-entitas dirancang menggunakan foreign key untuk menjaga integritas referensial data.',
    'Normal', doc)
make_para_before(ref, '4.4 Perancangan Interface (Mockup)', 'Heading 2', doc)
make_para_before(ref,
    'Perancangan antarmuka dilakukan sebelum implementasi untuk memastikan alur penggunaan yang intuitif. '
    'Mockup mencakup: halaman publik (papan informasi, login, registrasi), dashboard admin, '
    'dashboard instruktur, dan dashboard peserta. Desain menggunakan pendekatan mobile-first '
    'dengan Tailwind CSS untuk responsivitas lintas perangkat.',
    'Normal', doc)

# Rename existing "Tabel Database yang Digunakan" -> "4.5 Struktur Database"
tabel_db_para.clear()
tabel_db_para.text = '4.5 Struktur Database'
try:
    tabel_db_para.style = 'Heading 2'
except:
    pass

print("[INFO] Step 3: Change BAB V (Implementasi Sistem -> BAB V)")
# implementasi_para is currently Heading 2 inside BAB IV
# We need to insert a BAB V heading before it
ref5 = implementasi_para._element
make_para_before(ref5, 'BAB V', 'Heading 1', doc)
make_para_before(ref5, 'Implementasi dan Pengujian', 'Heading 1', doc)

# Rename "Implementasi Sistem" -> "5.1 Spesifikasi Teknologi"
implementasi_para.clear()
implementasi_para.text = '5.1 Spesifikasi Teknologi'
try:
    implementasi_para.style = 'Heading 2'
except:
    pass

# Rename "Pembahasan dan Hasil" -> "5.2 Implementasi Fitur"
pembahasan_hasil_para.clear()
pembahasan_hasil_para.text = '5.2 Implementasi Fitur'
try:
    pembahasan_hasil_para.style = 'Heading 2'
except:
    pass

print("[INFO] Step 4: Insert BAB II and BAB III")
# Add an explicit BAB II label before the existing company profile.
make_para_before(profil_para._element, 'BAB II', 'Heading 1', doc)
make_para_before(profil_para._element, 'Profil Perusahaan', 'Heading 1', doc)

# Insert BAB III before the existing theory/tools block so the company profile
# remains in BAB II and the tools are part of the PKL activities chapter.
ref3 = landasan_para._element

make_para_before(ref3, 'BAB III', 'Heading 1', doc)
make_para_before(ref3, 'Kegiatan Praktik Kerja Lapangan (PKL)', 'Heading 1', doc)

make_para_before(ref3, '3.1 Deskripsi Kegiatan PKL', 'Heading 2', doc)
make_para_before(ref3,
    'Praktik Kerja Lapangan (PKL) dilaksanakan di Direktorat Sistem dan Teknologi Informasi (DSTI) '
    'Universitas Pendidikan Indonesia (UPI). Selama PKL, penulis ditugaskan untuk merancang dan '
    'membangun ulang aplikasi Balai Kursus berbasis web menggunakan Laravel 10 sebagai framework '
    'utama. Kegiatan PKL mencakup tugas harian yang meliputi '
    'analisis kebutuhan, perancangan sistem, pengembangan fitur, pengujian, dan dokumentasi.',
    'Normal', doc)

make_para_before(ref3, '3.2 Tools dan Teknologi yang Digunakan', 'Heading 2', doc)
make_para_before(ref3,
    'Dalam pelaksanaan PKL, berbagai tools dan teknologi digunakan untuk mendukung proses '
    'pengembangan sistem, antara lain:\n'
    '- IDE & Editor: Visual Studio Code\n'
    '- Server Lokal: Laragon (Apache, PHP 8.1, MySQL)\n'
    '- Version Control: Git dan GitHub\n'
    '- Framework Backend: Laravel 10 (PHP)\n'
    '- Frontend: Tailwind CSS 3.1.0, Alpine.js 3.4.2, Vite 5.0.0\n'
    '- Database: MySQL dengan phpMyAdmin\n'
    '- Payment Gateway: Midtrans\n'
    '- Browser Pengujian: Google Chrome',
    'Normal', doc)

make_para_before(ref3, '3.3 Tugas Harian dan Output yang Dihasilkan', 'Heading 2', doc)
make_para_before(ref3,
    'Tugas harian selama PKL mencakup kegiatan teknis pengembangan dan koordinasi dengan '
    'pembimbing lapangan. Rangkuman tugas dan output yang dihasilkan adalah sebagai berikut:\n'
    '- Analisis sistem lama, studi kebutuhan, dan setup environment pengembangan.\n'
    '- Perancangan ERD dan struktur database, pembuatan migrasi, model, dan relasi Laravel.\n'
    '- Pengembangan autentikasi, manajemen peserta, instruktur, program, level, dan kursus.\n'
    '- Pengembangan penjadwalan, risalah, absensi, dan penilaian.\n'
    '- Integrasi pembayaran Midtrans dan penerbitan sertifikat digital dengan verifikasi publik.\n'
    '- Perbaikan bug, pengujian, dan penyusunan dokumentasi.\n'
    'Output utama adalah source code dan dokumentasi aplikasi Balai Kursus untuk Admin, '
    'Instruktur, dan Peserta.',
    'Normal', doc)

make_para_before(ref3, '3.4 Teknologi yang Dipelajari', 'Heading 2', doc)
make_para_before(ref3,
    'Selama PKL, penulis mempelajari dan mengaplikasikan berbagai teknologi baru, meliputi:\n'
    '- Arsitektur MVC Laravel 10 dengan fitur modern seperti Sanctum, Tinker, dan nWidart Modules.\n'
    '- Integrasi payment gateway Midtrans untuk transaksi online.\n'
    '- Implementasi Single Sign-On (SSO) menggunakan protokol CAS (subfission/cas).\n'
    '- Pembuatan sertifikat digital dengan QR code menggunakan endroid/qr-code dan barryvdh/laravel-dompdf.\n'
    '- Ekspor data ke Excel menggunakan Maatwebsite/Excel.\n'
    '- Pengembangan frontend responsif dengan Tailwind CSS dan Alpine.js.',
    'Normal', doc)

make_para_before(ref3, '3.5 Kendala dan Solusi', 'Heading 2', doc)
make_para_before(ref3,
    'Selama pelaksanaan PKL, terdapat beberapa kendala yang ditemui beserta solusi yang diterapkan:\n\n'
    'Kendala 1: Konflik jadwal pertemuan kursus pada saat penjadwalan manual.\n'
    'Solusi: Menambahkan validasi konflik jadwal pada backend (controller) sehingga sistem '
    'menolak input jadwal yang bertabrakan dengan jadwal yang sudah ada.\n\n'
    'Kendala 2: Integrasi SSO CAS UPI memerlukan konfigurasi khusus yang tidak terdokumentasi.\n'
    'Solusi: Berkoordinasi dengan tim DSTI untuk mendapatkan konfigurasi CAS yang benar dan '
    'melakukan pengujian bertahap.\n\n'
    'Kendala 3: Proses generate sertifikat PDF lambat untuk data dengan banyak gambar.\n'
    'Solusi: Menyederhanakan asset template, menguji proses generate pada data contoh, dan '
    'memisahkan proses preview dari publikasi sertifikat.',
    'Normal', doc)

print("[INFO] Step 5: Add BAB VI (Penutup) at end of document")
# Add BAB VI at end of body (before the final body section)
last_section = body[-1]  # w:sectPr

def append_heading(body, last_section, text, style_name, doc):
    new_p = OxmlElement('w:p')
    last_section.addprevious(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            p.text = text
            try:
                p.style = style_name
            except:
                pass
            return p
    return None

def append_para(body, last_section, text, style_name, doc):
    new_p = OxmlElement('w:p')
    last_section.addprevious(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            p.text = text
            try:
                p.style = style_name
            except:
                pass
            return p
    return None

# Add page break before BAB VI
br_p = OxmlElement('w:p')
br_r = OxmlElement('w:r')
br = OxmlElement('w:br')
br.set(qn('w:type'), 'page')
br_r.append(br)
br_p.append(br_r)
last_section.addprevious(br_p)

# Complete BAB V with the mandatory testing and deployment sections before
# starting BAB VI.
make_para_before(br_p,
    '5.3 Pengujian', 'Heading 2', doc)
make_para_before(br_p,
    'Pengujian dilakukan dengan pendekatan functional testing pada fitur utama. '
    'Ringkasan skenario dan hasilnya adalah sebagai berikut:\n'
    '- Autentikasi: login, registrasi, logout, verifikasi email, reset password, Google, dan CAS. Hasil: sesuai skenario.\n'
    '- Otorisasi: Admin, Instruktur, dan Peserta hanya mengakses menu sesuai role. Hasil: akses dibatasi middleware dan policy.\n'
    '- Pendaftaran dan placement: pendaftaran, score placement, pencocokan level, kuota, dan status pendaftaran. Hasil: sesuai alur aktif.\n'
    '- Pembayaran: pembuatan transaksi, callback berhasil/gagal, webhook, dan riwayat pembayaran. Hasil: status diperbarui sesuai respons gateway.\n'
    '- Jadwal: CRUD jadwal dan validasi bentrok lokasi/waktu. Hasil: jadwal bertabrakan ditolak.\n'
    '- Sertifikat: draft, preview, publish, revoke, restore draft, download, dan verifikasi publik. Hasil: sesuai hak akses dan status sertifikat.\n'
    'Pengujian perlu dijalankan kembali pada environment deployment sebelum digunakan oleh pengguna akhir.',
    'Normal', doc)
make_para_before(br_p,
    '5.4 Deployment', 'Heading 2', doc)
make_para_before(br_p,
    'Deployment aplikasi memerlukan server web Apache atau Nginx, PHP 8.1 atau lebih baru, '
    'MySQL, Composer, Node.js/NPM untuk asset frontend, konfigurasi environment, serta kredensial '
    'Midtrans dan CAS jika fitur tersebut diaktifkan. Tahapan deployment meliputi instalasi dependency, '
    'pengaturan .env, migrasi database, build asset Vite, konfigurasi storage, dan pengujian route. '
    'Dokumen yang dianalisis terutama menunjukkan environment lokal Laragon; status deployment produksi '
    'harus dikonfirmasi berdasarkan bukti server yang tersedia.',
    'Normal', doc)

append_heading(body, last_section, 'BAB VI', 'Heading 1', doc)
append_heading(body, last_section, 'Penutup', 'Heading 1', doc)

append_heading(body, last_section, '6.1 Kesimpulan', 'Heading 2', doc)
append_para(body, last_section,
    'Berdasarkan pelaksanaan Praktik Kerja Lapangan (PKL) di Direktorat Sistem dan Teknologi '
    'Informasi (DSTI) Universitas Pendidikan Indonesia (UPI), dapat ditarik kesimpulan sebagai berikut:\n\n'
    '1. Aplikasi Balai Kursus berbasis web berhasil dibangun ulang menggunakan Laravel 10 dengan '
    'arsitektur MVC yang terstruktur, menggantikan sistem lama yang belum terintegrasi.\n\n'
    '2. Sistem berhasil mengintegrasikan seluruh alur kerja pengelolaan kursus, mulai dari '
    'pendaftaran online, penjadwalan, pencatatan absensi dan risalah, penilaian peserta, '
    'pembayaran melalui Midtrans, hingga penerbitan sertifikat digital dengan QR code.\n\n'
    '3. Tiga peran pengguna (Admin, Instruktur, Peserta) telah diimplementasikan dengan hak akses '
    'yang berbeda sesuai kebutuhan operasional DSTI UPI.\n\n'
    '4. Melalui PKL ini, penulis memperoleh pengalaman nyata dalam pengembangan perangkat lunak '
    'profesional, termasuk penerapan best practice Laravel, integrasi third-party API, '
    'dan manajemen proyek berbasis Git.',
    'Normal', doc)

append_heading(body, last_section, '6.2 Saran', 'Heading 2', doc)
append_para(body, last_section,
    'Demi pengembangan sistem yang lebih baik ke depannya, penulis menyampaikan beberapa saran:\n\n'
    '1. Notifikasi Real-time: Menambahkan fitur notifikasi (email/push notification) untuk '
    'menginformasikan jadwal kursus, status pembayaran, dan penerbitan sertifikat kepada peserta.\n\n'
    '2. Fitur Pelaporan Lanjutan: Menambahkan modul laporan yang lebih komprehensif dengan '
    'visualisasi data (grafik, chart) untuk memudahkan analisis kinerja kursus oleh manajemen.\n\n'
    '3. Aplikasi Mobile: Mengembangkan versi mobile (Android/iOS) atau Progressive Web App (PWA) '
    'agar peserta dan instruktur dapat mengakses sistem dengan lebih mudah dari perangkat mobile.\n\n'
    '4. Keamanan Sistem: Melakukan audit keamanan secara berkala, menambahkan 2-Factor Authentication '
    '(2FA), dan memastikan seluruh data sensitif terenkripsi dengan baik.\n\n'
    '5. Pemeliharaan Berkala: Melakukan pembaruan framework dan library secara rutin untuk '
    'menjaga keamanan dan performa sistem jangka panjang.',
    'Normal', doc)

print("[INFO] Saving output...")
doc.save(str(output))
print(f"[DONE] Saved to: {output}")
