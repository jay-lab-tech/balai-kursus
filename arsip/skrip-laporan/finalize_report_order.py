from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.oxml.ns import qn

path = Path('docs/LAPORAN KEGIATAN PKL RPL (UPI) - FINAL.docx')
doc = Document(str(path))

def wrap(el):
    return next(p for p in doc.paragraphs if p._element is el)

def insert_after(ref, text, style='Normal'):
    new = OxmlElement('w:p')
    ref._element.addnext(new)
    p = wrap(new)
    p.style = style
    p.add_run(text)
    return p

# Make the mandatory BAB III result heading explicit.
for p in doc.paragraphs:
    if p.text.strip() == '3.3 Tugas Harian dan Output yang Dihasilkan':
        p.text = '3.3 Hasil Kegiatan, Tugas Harian, dan Output'
        p.style = 'Heading 2'
        break

# Expand the ERD section with the complete active relationship map.
erd = next((p for p in doc.paragraphs if p.text.strip() == '4.3 Entity Relationship Diagram (ERD)'), None)
if erd is not None and not any('Relasi aktif ERD' in p.text for p in doc.paragraphs):
    insert_after(erd,
        'Relasi aktif ERD: users memiliki satu profil pesertas atau instrukturs; pesertas memiliki banyak pendaftarans; '
        'programs memiliki banyak kursuses dan pendaftarans; levels memiliki banyak kursuses, pendaftarans, serta '
        'assignment peserta/instruktur; kursuses memiliki banyak jadwals, risalahs, pendaftarans, dan assignment; '
        'jadwals terhubung ke lokasis, kelas, hari, dan pembuat; pendaftarans memiliki payments, scores, dan absensis; '
        'risalahs memiliki absensis; certificate_templates memiliki certificates; certificates terhubung ke kursus, '
        'peserta, dan user. Assignment aktif peserta menggunakan peserta_kursus_levels dan assignment instruktur '
        'menggunakan instruktur_kursus_levels. Tabel peserta_kursus dan field instruktur_id_2 diberi status legacy.',
        'Normal')

# Put Daftar Pustaka and Lampiran after BAB VI, which is the conventional report order.
start = next((p for p in doc.paragraphs if p.text.strip() == 'Daftar Pustaka'), None)
chapter6 = next((p for p in doc.paragraphs if p.text.strip() == 'BAB VI'), None)
if start is not None and chapter6 is not None and start._element.getparent() is chapter6._element.getparent():
    body = doc.element.body
    sect = body.sectPr
    moving = []
    node = start._element
    while node is not None and node is not chapter6._element:
        nxt = node.getnext()
        moving.append(node)
        node = nxt
    for el in moving:
        body.remove(el)
    for el in moving:
        sect.addprevious(el)

# Ensure all moved paragraphs remain 12 pt Times New Roman and 1.5 spaced.
for p in doc.paragraphs:
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Times New Roman')
        r._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        if p.style.name.lower().startswith('heading') or p.text.strip().startswith('BAB '):
            r.font.size = Pt(14)
            r.bold = True
        else:
            r.font.size = Pt(12)

doc.save(str(path))
print(f'Finalized {path}')
