from docx import Document
from pathlib import Path

p = Path('docs/LAPORAN KEGIATAN PKL RPL (UPI) - RAPI.docx')
doc = Document(str(p))
print('Total paragraphs:', len(doc.paragraphs))
print('=' * 70)
for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    style = para.style.name if para.style else '?'
    # Print headings and any line mentioning BAB
    if style.startswith('Heading') or style.startswith('Title') or txt.upper().startswith('BAB'):
        print(f'[{i}] ({style}) {txt}')
