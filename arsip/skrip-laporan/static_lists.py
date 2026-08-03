from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

path = Path('docs/LAPORAN KEGIATAN PKL RPL (UPI) - FINAL.docx')
doc = Document(str(path))

table_items = [p.text.strip() for p in doc.paragraphs if p.text.strip().startswith('Table ')]
figure_items = [p.text.strip() for p in doc.paragraphs if p.text.strip().startswith('Figure ')]

def replace_next_empty(title, items):
    ps = doc.paragraphs
    for i, p in enumerate(ps):
        if p.text.strip().upper() == title:
            for nxt in ps[i + 1:i + 4]:
                if not nxt.text.strip():
                    nxt.text = '\n'.join(items)
                    nxt.style = 'Normal'
                    nxt.paragraph_format.line_spacing = 1.5
                    for r in nxt.runs:
                        r.font.name = 'Times New Roman'
                        r._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Times New Roman')
                        r._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                        r.font.size = Pt(12)
                    return

replace_next_empty('DAFTAR TABEL', table_items)
replace_next_empty('DAFTAR GAMBAR', figure_items)
doc.save(str(path))
print(f'Static lists written: {len(table_items)} tables, {len(figure_items)} figures')
