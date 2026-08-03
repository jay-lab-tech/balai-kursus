from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

SRC = Path('docs/LAPORAN PKL RPL - VERSI SIAP 2.docx')
OUT = Path('docs/arsip-laporan/LAPORAN PKL RPL - VERSI NATURAL.docx')
doc = Document(str(SRC))

replacements = [
    ('Penulis mengucapkan terima kasih kepada pihak sekolah, pembimbing, keluarga, dan semua pihak yang telah membantu selama kegiatan PKL. Penulis menyadari masih terdapat kekurangan dalam laporan ini, sehingga saran dan kritik sangat diharapkan.',
     'Penulis mengucapkan terima kasih kepada pihak sekolah, pembimbing, keluarga, dan semua pihak yang telah membantu selama PKL. Laporan ini masih memiliki kekurangan, sehingga kritik dan saran tetap diperlukan untuk memperbaikinya.'),
    ('Tujuan kegiatan ini adalah membangun aplikasi administrasi kursus berbasis web, menerapkan pengetahuan yang diperoleh di sekolah, dan memahami proses pengembangan perangkat lunak dalam lingkungan kerja.',
     'Kegiatan ini bertujuan membangun aplikasi administrasi kursus berbasis web, menerapkan pengetahuan dari sekolah, dan memahami proses pengembangan perangkat lunak di lingkungan kerja.'),
    ('Kegiatan PKL berfokus pada pembangunan ulang aplikasi Balai Kursus berbasis web. Kegiatan dilakukan melalui tahapan analisis kebutuhan, perancangan database dan interface, implementasi modul, integrasi layanan, pengujian, perbaikan, dan penyusunan dokumentasi.',
     'Selama PKL, penulis melanjutkan coding aplikasi Balai Kursus. Pekerjaan meliputi analisis kebutuhan, perancangan database dan interface, pengembangan modul, integrasi layanan, perbaikan bug, pengujian, dan dokumentasi.'),
    ('Rangkaian kegiatan berikut disusun berdasarkan tahapan yang terlihat pada project dan dokumentasi teknis. Tanggal dapat disesuaikan kembali dengan jurnal asli sebelum dicetak.',
     'Pembagian kegiatan di bawah diringkas dari pekerjaan coding dan dokumentasi project. Tanggalnya dapat disesuaikan dengan jurnal asli sebelum laporan dicetak.'),
    ('Masalah utama yang menjadi dasar pembangunan ulang adalah kebutuhan terhadap sistem yang menghubungkan administrasi peserta, program, level, kursus, jadwal, pembelajaran, pembayaran, nilai, dan sertifikat. Sistem juga harus membedakan kewenangan tiap role dan menyediakan data yang dapat ditelusuri dari satu proses ke proses berikutnya.',
     'Project perlu menghubungkan administrasi peserta, program, level, kursus, jadwal, pembelajaran, pembayaran, nilai, dan sertifikat dalam satu alur. Akses setiap role juga perlu dibatasi agar data hanya dapat dikelola oleh pengguna yang berwenang.'),
    ('ERD digunakan untuk melihat hubungan antar-entitas utama dalam sistem. Tabel database tidak ditulis satu per satu karena rincian field sudah tersedia pada migration dan dokumentasi teknis project. Diagram berikut merangkum entitas akun, profil pengguna, katalog kursus, pendaftaran, pembayaran, pembelajaran, penilaian, jadwal, dan sertifikat.',
     'ERD memperlihatkan hubungan antar-entitas utama dalam sistem. Rincian field tetap berada pada migration dan dokumentasi teknis project, sedangkan laporan ini menampilkan hubungan akun, profil pengguna, katalog kursus, pendaftaran, pembayaran, pembelajaran, penilaian, jadwal, dan sertifikat.'),
    ('Entitas utama yang ditampilkan pada diagram adalah users, pesertas, instrukturs, programs, levels, kursuses, pendaftarans, payments, jadwals, risalahs, absensis, scores, dan certificates. Entitas pendukung seperti lokasi, kelas, hari, template sertifikat, serta tabel penghubung digunakan untuk melengkapi proses bisnis dan ditunjukkan pada relasi/struktur project.',
     'Diagram memuat users, pesertas, instrukturs, programs, levels, kursuses, pendaftarans, payments, jadwals, risalahs, absensis, scores, dan certificates. Lokasi, kelas, hari, template sertifikat, serta tabel penghubung menjadi entitas pendukung dalam proses bisnis project.'),
    ('Perancangan interface mengikuti pembagian role dan alur bisnis. Halaman publik dan autentikasi menjadi pintu masuk, dashboard menjadi pusat navigasi, sedangkan halaman CRUD memakai pola form, validasi, tabel, pencarian, dan pagination. Interface dirancang responsif menggunakan Tailwind CSS.',
     'Interface disusun sesuai role dan alur bisnis. Halaman publik dan autentikasi menjadi akses awal, dashboard menjadi pusat navigasi, sedangkan halaman CRUD memakai form, validasi, tabel, pencarian, dan pagination. Tampilan dibuat responsif dengan Tailwind CSS.'),
    ('Visual berikut diambil dari project yang dijalankan pada environment lokal. Gambar ini memperlihatkan bentuk halaman, warna, navigasi, dan komponen interface yang digunakan sebagai dasar pembahasan perancangan.',
     'Gambar pada bagian ini diambil dari project yang dijalankan di environment lokal. Tampilan tersebut memperlihatkan halaman, warna, navigasi, dan komponen interface yang dibahas pada bagian perancangan.'),
    ('Project menggunakan PHP 8.1+, Laravel 10.10, MySQL, Laravel Modules, Blade, Vite, Tailwind CSS, Alpine.js, Midtrans, CAS/SSO, Dompdf, QR Code, dan Maatwebsite Excel. Laragon digunakan sebagai environment lokal. Credential layanan tidak dicantumkan dalam laporan.',
     'Project dibangun dengan PHP 8.1+, Laravel 10.10, MySQL, Laravel Modules, Blade, Vite, Tailwind CSS, Alpine.js, Midtrans, CAS/SSO, Dompdf, QR Code, dan Maatwebsite Excel. Laragon dipakai untuk menjalankan project secara lokal. Credential layanan tidak dicantumkan dalam laporan.'),
    ('Implementasi fitur dikelompokkan berdasarkan role dan proses utamanya agar pembahasan tetap ringkas. Screenshot berikut dipilih untuk mewakili fitur, sedangkan keseluruhan dokumentasi tampilan dicantumkan pada lampiran.',
     'Fitur dibahas berdasarkan role dan proses utama agar isinya tetap ringkas. Screenshot yang paling mewakili setiap fitur ditampilkan pada BAB V, sedangkan gambar tambahan ditempatkan pada lampiran.'),
    ('Pengujian yang dapat dilakukan pada environment saat penyusunan laporan adalah pemeriksaan tampilan melalui browser lokal dan pemeriksaan screenshot fitur. Feature test Laravel sudah tersedia pada folder tests, tetapi belum dapat dijalankan sampai selesai karena service MySQL pada 127.0.0.1:3306 tidak aktif. Pengujian database dan pengujian online harus diulang setelah database serta server tersedia.',
     'Saat laporan disusun, penulis menguji tampilan melalui browser lokal dan memeriksa screenshot fitur. Feature test Laravel sudah tersedia di folder tests, tetapi belum selesai dijalankan karena service MySQL pada 127.0.0.1:3306 tidak aktif. Pengujian database dan pengujian online perlu diulang setelah database serta server tersedia.'),
    ('Deployment nyata belum diisi pada versi ini. Bagian berikut disiapkan untuk dilengkapi setelah informasi server tersedia:',
     'Project belum dideploy. Data URL, server, database produksi, dan hasil pengujian online akan ditambahkan setelah deployment dilakukan:'),
    ('Berdasarkan hasil pengembangan, aplikasi Balai Kursus telah memiliki alur utama untuk mengelola program, pendaftaran, placement, kelas, pembayaran, jadwal, risalah, absensi, nilai, dan sertifikat. Pengembangan ini juga memberi pengalaman kepada penulis dalam menggunakan Laravel, membuat relasi database, menghubungkan layanan eksternal, dan melakukan pengujian.',
     'Project Balai Kursus sudah memiliki alur untuk mengelola program, pendaftaran, placement, kelas, pembayaran, jadwal, risalah, absensi, nilai, dan sertifikat. Selama mengerjakannya, penulis mempraktikkan penggunaan Laravel, relasi database, integrasi layanan eksternal, dan pengujian aplikasi.'),
    ('Screenshot berikut menjadi bukti visual dari project yang dijalankan pada environment lokal.',
     'Screenshot di lampiran menjadi bukti visual project yang dijalankan di environment lokal.'),
    ('Diagram dan tampilan berikut merupakan visual yang digunakan dalam analisis dan perancangan project.',
     'Lampiran ini memuat diagram dan tampilan yang dipakai pada analisis serta perancangan project.'),
    ('Kegiatan PKL dilakukan dengan melanjutkan coding project Balai Kursus secara berkelompok dengan satu orang rekan. Pembagian pekerjaan mengikuti kesepakatan tim; penulis mengerjakan bagian coding lanjutan, integrasi fitur, pemeriksaan tampilan, dan dokumentasi, sedangkan rekan mengerjakan bagian fitur lain sesuai pembagian tugas dan melakukan koordinasi perubahan project.',
     'Selama PKL, penulis melanjutkan coding project Balai Kursus bersama satu orang rekan. Pembagian pekerjaan mengikuti kesepakatan tim. Penulis mengerjakan coding lanjutan, integrasi fitur, pemeriksaan tampilan, dan dokumentasi, sedangkan rekan mengerjakan bagian fitur lain serta membantu meninjau perubahan project.'),
    ('Universitas Pendidikan Indonesia (UPI) merupakan salah satu Perguruan Tinggi Negeri Badan Hukum (PTN-BH) di Indonesia yang berfokus pada bidang kependidikan dan non-kependidikan. Perjalanan sejarah berdirinya UPI dapat diklasifikasikan ke dalam beberapa fase perkembangan sebagai berikut:',
     'Universitas Pendidikan Indonesia (UPI) merupakan Perguruan Tinggi Negeri Badan Hukum (PTN-BH) yang bergerak di bidang kependidikan dan non-kependidikan. Sejarah UPI dibagi ke dalam beberapa fase perkembangan:'),
    ('Guna mewujudkan visi tersebut, misi yang diemban oleh Universitas Pendidikan Indonesia adalah sebagai berikut:',
     'Untuk mencapai visi tersebut, Universitas Pendidikan Indonesia menjalankan misi berikut:'),
    ('Universitas Pendidikan Indonesia memiliki lambang resmi berbentuk bulat dengan perpaduan warna merah dan putih, serta komponen visual yang khas. Setiap unsur yang terkandung di dalam lambang tersebut memiliki arti dan makna filosofis sebagai berikut:',
     'Lambang Universitas Pendidikan Indonesia berbentuk bulat dengan perpaduan warna merah dan putih. Setiap unsurnya memiliki makna filosofis:'),
    ('Secara operasional, DSTI UPI menjalankan fungsi-fungsi utama sebagai berikut:',
     'Dalam kegiatan operasionalnya, DSTI UPI menjalankan fungsi utama berikut:'),
]

def set_text(p, text):
    p.text = text
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Times New Roman')
        r._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 0)

for p in doc.paragraphs:
    for old, new in replacements:
        if p.text.strip() == old:
            set_text(p, new)
            break
    if 'berikut' in p.text.lower():
        set_text(p, p.text.replace('berikut', 'tersebut').replace('Berikut', 'Tersebut'))

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                text = p.text
                text = text.replace('Hasil yang Diharapkan', 'Target pengujian')
                text = text.replace('Visualisasi Interface dari Project', 'Tampilan Interface Project')
                text = text.replace('berikut', 'tersebut').replace('Berikut', 'Tersebut')
                if text != p.text:
                    set_text(p, text)

doc.save(str(OUT))
print(f'Saved {OUT}')
