# -*- coding: utf-8 -*-
"""Sisipkan screenshot ke 'LAPORAN KEGIATAN PKL RPL (UPI).docx'
menggantikan paragraf placeholder '[ Sisipkan tangkapan layar di sini ]'.
Gambar dicocokkan via caption 'Figure 4.x <deskripsi>' dengan manifest.json.
"""
import json, re
from docx import Document
from docx.shared import Emu

PATH = 'docs/LAPORAN KEGIATAN PKL RPL (UPI).docx'
PLACEHOLDER = '[ Sisipkan tangkapan layar di sini ]'

manifest = json.load(open('docs/screenshots/manifest.json', encoding='utf-8'))
doc = Document(PATH)

# lebar konten (lebar halaman dikurangi margin kiri-kanan)
sec = doc.sections[0]
content_w = sec.page_width - sec.left_margin - sec.right_margin

paras = doc.paragraphs
inserted, missing = [], []

def desc_from_caption(t):
    return re.sub(r'^Figure\s+4\.\d+\s+', '', t).strip()

for i, p in enumerate(paras):
    if p.text.strip() != PLACEHOLDER:
        continue
    # cari caption pada paragraf berikutnya
    desc = None
    for j in range(i + 1, min(i + 4, len(paras))):
        t = paras[j].text.strip()
        if t.startswith('Figure 4.'):
            desc = desc_from_caption(t)
            break
    img = manifest.get(desc) if desc else None
    if not img:
        missing.append(desc or '(caption tidak ditemukan)')
        continue
    # bersihkan run placeholder, lalu sisipkan gambar
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run()
    run.add_picture(img, width=content_w)
    inserted.append(desc)

doc.save(PATH)
print('Disisipkan:', len(inserted))
for d in inserted:
    print('  +', d)
print('Tetap placeholder:', len(missing))
for d in missing:
    print('  -', d)
