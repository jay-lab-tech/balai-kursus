from pathlib import Path
from zipfile import ZipFile
import tempfile
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('docs/LAPORAN PKL RPL - VERSI SIAP 2.docx')
IMG = Path('docs/screenshots')
SCHOOL_LOGO = Path(r'C:\Users\ASUS\AppData\Local\Temp\balai-old-docx-media\x\word\media\image1.png')
OLD_REPORT = Path('docs/LAPORAN KEGIATAN PKL RPL (UPI) - RAPI.docx')

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(3)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(4)
section.right_margin = Cm(2.5)

def set_font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic

def style_para(p, align=None, before=0, after=0, first_indent=True):
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1)
    if align is not None:
        p.alignment = align
    for r in p.runs:
        set_font(r)

def para(text='', align=None, before=0, after=0, first_indent=True):
    p = doc.add_paragraph()
    if text:
        p.add_run(text)
    style_para(p, align, before, after, first_indent)
    return p

def heading(text, level=1):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, size=14, bold=True)
    return p

def chapter(number, title):
    doc.add_page_break()
    p = heading(f'BAB {number}', 1)
    p.paragraph_format.page_break_before = False
    heading(title, 1)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    style_para(p, first_indent=False)
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    style_para(p, first_indent=False)
    return p

def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            for r in p.runs:
                set_font(r, size=11, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_font(r, size=11)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table

def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)

def add_wireframe(title, rows):
    heading(title, 3)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.rows[0].cells[0].merge(table.rows[0].cells[1])
    top = table.cell(0, 0)
    top.text = 'WIREFRAME / RANCANGAN UI'
    shade_cell(top, 'D9EAF7')
    for p in top.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        for r in p.runs: set_font(r, size=11, bold=True)
    for label, content in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = content
        shade_cell(cells[0], 'EEF4F8')
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs: set_font(r, size=11, bold=(cell is cells[0]))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table

def add_screenshot(filename, caption, width=13.5):
    path = IMG / filename
    if not path.exists():
        para(f'[Screenshot tidak ditemukan: {filename}]', first_indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    style_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    c = para(caption, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    for r in c.runs:
        set_font(r, size=11, bold=True)

def add_page_field(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    r = p.add_run('Halaman ')
    set_font(r, size=10)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    p._p.append(fld)

def extract_old_image(member):
    target = Path(tempfile.gettempdir()) / 'balai-kursus-old-media' / Path(member).name
    target.parent.mkdir(parents=True, exist_ok=True)
    if not target.exists():
        with ZipFile(OLD_REPORT) as z:
            target.write_bytes(z.read('word/' + member))
    return target

def add_old_para(old_doc, index, heading_level=None):
    text = old_doc.paragraphs[index].text.strip()
    if not text:
        return
    if heading_level:
        heading(text, heading_level)
    else:
        para(text)

for style_name in ['Normal', 'List Bullet', 'List Number']:
    try:
        st = doc.styles[style_name]
        st.font.name = 'Times New Roman'
        st._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Times New Roman')
        st._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        st.font.size = Pt(12)
        st.paragraph_format.line_spacing = 1.5
    except KeyError:
        pass
for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    st = doc.styles[style_name]
    st.font.name = 'Times New Roman'
    st._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Times New Roman')
    st._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    st.font.size = Pt(14)
    st.font.bold = True
    st.paragraph_format.line_spacing = 1.5

# Sampul mengikuti susunan file laporan sebelumnya: judul terpusat,
# logo sekolah di tengah, identitas penulis, lalu identitas sekolah.
for _ in range(3): para('', first_indent=False)
for text in ['LAPORAN KEGIATAN',
             'PEMBANGUNAN ULANG APLIKASI BALAI KURSUS',
             'BERBASIS WEB MENGGUNAKAN LARAVEL 10',
             'UNIVERSITAS PENDIDIKAN INDONESIA']:
    p = para(text, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    for r in p.runs: set_font(r, size=14, bold=True)
for _ in range(5): para('', first_indent=False)
if SCHOOL_LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(SCHOOL_LOGO), width=Cm(4.2))
    style_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
for _ in range(5): para('', first_indent=False)
para('Oleh:', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
names = doc.add_table(rows=2, cols=2)
names.alignment = WD_TABLE_ALIGNMENT.CENTER
names.style = 'Table Grid'
for row, values in zip(names.rows, [
    ['Galih Aziz Wahyudi', '0088804868'],
    ['Azhar Noermansyah', '0097477130'],
]):
    for cell, value in zip(row.cells, values):
        cell.text = value
for row in names.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: set_font(r, size=12, bold=False)
para('SMKN 1 SUBANG', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
para('Konsentrasi Keahlian Rekayasa Perangkat Lunak (RPL)', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
para('Jalan Arief Rahman Hakim No. 35, Cigadung, Kecamatan Subang', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
para('Kabupaten Subang, Jawa Barat 41213', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
para('2026', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

# Lembar pengesahan mengikuti dua halaman pengesahan pada laporan sebelumnya.
# Section break ditambahkan pada tahap penomoran halaman; page break tambahan
# di sini akan menghasilkan satu halaman kosong sebelum pengesahan.
heading('Lembar Pengesahan Sekolah', 1)
for text in ['PEMBANGUNAN ULANG APLIKASI BALAI KURSUS BERBASIS WEB MENGGUNAKAN',
             'LARAVEL 10', 'UNIVERSITAS PENDIDIKAN INDONESIA']:
    p = para(text, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    for r in p.runs: set_font(r, size=12, bold=True)
para('Telah disahkan pada:', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
para('Hari :', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
para('Tanggal :', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
add_table(['Kepala Konsentrasi Keahlian RPL', 'Mengetahui', 'Guru Pembimbing PKL'], [[
    'Rizal Suyaman, S.Kom.\nNIP 198601262022211007',
    'Menyetujui,\nKepala Sekolah SMKN 1 Subang',
    'Yani Yulyanti, S.T.\nNIP 198812132020122010'
]], [4.6, 4.6, 4.6])
para('[Tempat tanda tangan, nama terang, dan stempel sekolah diisi kemudian]', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
doc.add_page_break()
heading('Lembar Pengesahan Perusahaan (IDUKA)', 1)
for text in ['PEMBANGUNAN ULANG APLIKASI BALAI KURSUS BERBASIS WEB MENGGUNAKAN',
             'LARAVEL 10', 'UNIVERSITAS PENDIDIKAN INDONESIA']:
    p = para(text, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    for r in p.runs: set_font(r, size=12, bold=True)
para('Telah disahkan pada:', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
para('Hari :', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
para('Tanggal :', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
add_table(['Kepala DSTI UPI', '', 'Pembimbing PKL DSTI UPI'], [[
    'Dr. Cepi Riyana, M.Pd.\nNIP 197512302001121001',
    'Menyetujui',
    'Aceng Sobana, S.Pd., M.T.\nNIP 197803112014091001'
]], [4.6, 4.6, 4.6])
para('[Tempat tanda tangan, nama terang, dan stempel IDUKA diisi kemudian]', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

# Kata pengantar
doc.add_page_break()
heading('KATA PENGANTAR', 1)
para('Puji syukur penulis panjatkan kepada Tuhan Yang Maha Esa karena laporan kegiatan Praktik Kerja Lapangan ini dapat diselesaikan. Laporan ini disusun sebagai bentuk pertanggungjawaban atas kegiatan PKL sekaligus dokumentasi proses pembangunan ulang aplikasi Balai Kursus di Direktorat Sistem dan Teknologi Informasi Universitas Pendidikan Indonesia.')
para('Penulis menyampaikan terima kasih kepada pihak sekolah, pembimbing, keluarga, dan seluruh pihak yang telah memberikan arahan serta dukungan selama pelaksanaan PKL. Penulis menyadari bahwa laporan ini masih dapat disempurnakan, sehingga kritik dan saran yang membangun sangat diharapkan.')
para('Bandung, Juni 2026', align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)
para('Penulis', align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)

# Daftar isi and lists
doc.add_page_break()
heading('DAFTAR ISI', 1)
p = para('', first_indent=False)
fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u'); p._p.append(fld)
heading('DAFTAR TABEL', 1)
para('Tabel 3.1 Tools dan teknologi\nTabel 3.2 Tugas harian dan output\nTabel 3.3 Kendala dan solusi\nTabel 4.1 Kebutuhan sistem\nTabel 4.2 Ringkasan entitas dan relasi ERD\nTabel 4.3 Komponen perancangan interface\nTabel 5.1 Skenario dan hasil pengujian\nTabel 5.2 Kebutuhan deployment\nTabel E.1 Ringkasan jurnal kegiatan', first_indent=False)
heading('DAFTAR GAMBAR', 1)
para('Gambar 4.1 ERD konseptual\nGambar 4.2–4.3 Screenshot rancangan interface\nGambar 4.4–4.6 Mockup visual UI project\nGambar 5.1–5.8 Screenshot implementasi fitur', first_indent=False)

# BAB I
chapter('I', 'PENDAHULUAN')
heading('1.1 Latar Belakang', 2)
para('Perkembangan teknologi informasi mendorong lembaga pendidikan nonformal untuk menyediakan layanan yang cepat, terintegrasi, dan mudah diakses. Pada pengelolaan kursus, proses pendaftaran, penempatan peserta, penjadwalan, absensi, penilaian, pembayaran, dan penerbitan sertifikat perlu dikelola dalam satu sistem agar data tidak tersebar dan pekerjaan administrasi dapat dilakukan secara lebih efisien.')
para('Aplikasi Balai Kursus dikembangkan untuk mendukung kebutuhan tersebut. Sistem menyediakan akses berbeda bagi Admin, Instruktur, dan Peserta, serta menghubungkan proses pendaftaran dengan placement, kelas, pembayaran Midtrans, kegiatan pembelajaran, penilaian, dan sertifikat.')
heading('1.2 Identifikasi Masalah', 2)
numbered('Proses administrasi kursus perlu dikelola secara terintegrasi dari pendaftaran sampai sertifikat.')
numbered('Data peserta, instruktur, program, level, kursus, jadwal, absensi, dan nilai membutuhkan struktur yang konsisten.')
numbered('Penempatan peserta perlu mempertimbangkan hasil placement dan kuota kelas.')
numbered('Status pembayaran dan sertifikat perlu dapat dipantau secara digital.')
heading('1.3 Batasan Masalah', 2)
bullet('Sistem ditujukan untuk satu instansi dan menggunakan tiga role: Admin, Instruktur, dan Peserta.')
bullet('Pembayaran daring dibatasi pada integrasi Midtrans dan pemutakhiran status pembayaran; akuntansi di luar ruang lingkup.')
bullet('Sistem bukan LMS penuh; fitur pembelajaran dibatasi pada jadwal, risalah, absensi, dan penilaian.')
bullet('Deployment produksi tidak dibahas sebagai instalasi server yang telah diverifikasi; laporan menjelaskan kebutuhan dan tahapan deployment.')
heading('1.4 Tujuan', 2)
para('Tujuan kegiatan ini adalah menerapkan pengetahuan rekayasa perangkat lunak dalam pembangunan ulang aplikasi Balai Kursus, menghasilkan sistem administrasi kursus berbasis web, serta memperoleh pengalaman dalam analisis, perancangan, implementasi, pengujian, dan dokumentasi aplikasi.')
heading('1.5 Manfaat', 2)
bullet('Bagi instansi: membantu pengelolaan data dan proses kursus secara terintegrasi.')
bullet('Bagi pengguna: menyediakan akses layanan sesuai role melalui browser.')
bullet('Bagi penulis: memperkuat pengalaman pengembangan aplikasi Laravel dan integrasi layanan pihak ketiga.')

# BAB II
chapter('II', 'PROFIL PERUSAHAAN')
old = Document(str(OLD_REPORT))
heading('2.1 Sejarah Universitas Pendidikan Indonesia', 2)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(extract_old_image('media/image2.jpeg')), width=Cm(12.5)); style_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
c = para('Gambar 2.1 Sejarah Universitas Pendidikan Indonesia', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
for r in c.runs: set_font(r, size=11, bold=True)
for idx in [102, 103, 104, 105, 106, 107, 108]:
    add_old_para(old, idx, 3 if idx in [103, 105, 107] else None)
heading('2.2 Visi, Misi, dan Tujuan UPI', 2)
add_old_para(old, 110)
add_old_para(old, 111, 3); add_old_para(old, 112)
add_old_para(old, 113, 3); add_old_para(old, 114)
for idx in [115, 116, 117, 118]: add_old_para(old, idx)
add_old_para(old, 119, 3); add_old_para(old, 120)
for idx in [121, 122, 123, 124]: add_old_para(old, idx)
heading('2.3 Arti dan Makna Simbol UPI', 2)
add_old_para(old, 128)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(extract_old_image('media/image3.jpeg')), width=Cm(7.0)); style_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
c = para('Gambar 2.2 Logo Universitas Pendidikan Indonesia', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
for r in c.runs: set_font(r, size=11, bold=True)
for idx in [130, 131, 132, 133, 134, 135, 136, 137, 138, 139, 140, 141, 142, 143, 144]:
    add_old_para(old, idx, 3 if idx in [130, 132, 137, 139, 141] else None)
heading('2.4 Profil Perusahaan', 2)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(extract_old_image('media/image4.jpeg')), width=Cm(13.0)); style_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
c = para('Gambar 2.3 Direktorat Sistem dan Teknologi Informasi', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
for r in c.runs: set_font(r, size=11, bold=True)
for idx in [148, 149, 150, 151, 152, 153, 154]: add_old_para(old, idx)
heading('2.5 Struktur Organisasi DSTI', 2)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(extract_old_image('media/image5.png')), width=Cm(13.0)); style_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
c = para('Gambar 2.4 Struktur Organisasi DSTI', align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
for r in c.runs: set_font(r, size=11, bold=True)
heading('2.6 Profil DSTI UPI', 2)
for idx in [158, 159]: add_old_para(old, idx)
add_old_para(old, 160, 3); add_old_para(old, 161)
add_old_para(old, 167, 3)
for idx in [168, 169, 170, 171]: add_old_para(old, idx)
add_old_para(old, 172, 3); add_old_para(old, 173)
for idx in [174, 176, 178]: add_old_para(old, idx)

# BAB III
chapter('III', 'KEGIATAN PKL')
heading('3.1 Deskripsi Kegiatan', 2)
para('Kegiatan PKL berfokus pada pembangunan ulang aplikasi Balai Kursus berbasis web. Kegiatan dilakukan melalui tahapan analisis kebutuhan, perancangan database dan interface, implementasi modul, integrasi layanan, pengujian, perbaikan, dan penyusunan dokumentasi.')
heading('3.2 Tools dan Teknologi', 2)
add_table(['Kelompok', 'Tools/Teknologi'], [
    ['Editor dan server lokal', 'Visual Studio Code, Laragon, Apache, PHP 8.1, MySQL'],
    ['Backend', 'Laravel 10, Eloquent ORM, Laravel Breeze, Sanctum'],
    ['Frontend', 'Blade, Tailwind CSS, Alpine.js, Vite, Axios'],
    ['Layanan pihak ketiga', 'Midtrans, Google OAuth, CAS/SSO'],
    ['Dokumentasi dan kontrol versi', 'Git, GitHub, phpMyAdmin, Chrome DevTools'],
], [5.0, 9.0])
heading('3.3 Tugas Harian dan Output', 2)
para('Rangkaian kegiatan berikut disusun berdasarkan tahapan yang terlihat pada project dan dokumentasi teknis. Tanggal dapat disesuaikan kembali dengan jurnal asli sebelum dicetak.', first_indent=False)
add_table(['Periode/Tanggal', 'Tugas yang Dilakukan', 'Output'], [
    ['Minggu 1', 'Mempelajari alur bisnis Balai Kursus, role pengguna, dan kebutuhan sistem.', 'Catatan analisis kebutuhan dan batasan sistem.'],
    ['Minggu 2', 'Menganalisis struktur project serta merancang relasi data program, level, kursus, peserta, dan pendaftaran.', 'Rancangan ERD dan pemetaan migration/model.'],
    ['Minggu 3', 'Mengerjakan autentikasi, pembagian akses, dashboard, dan halaman data master.', 'Fitur login, role, dashboard, serta CRUD data utama.'],
    ['Minggu 4', 'Mengerjakan proses placement, kursus, jadwal, risalah, absensi, dan penilaian.', 'Alur kelas dan pembelajaran dapat digunakan pada environment lokal.'],
    ['Minggu 5', 'Menghubungkan pembayaran Midtrans, sertifikat, QR verifikasi, dan ekspor laporan.', 'Integrasi layanan dan fitur sertifikat digital.'],
    ['Minggu 6', 'Melakukan pengujian, memperbaiki kendala, mengambil dokumentasi tampilan, dan menyusun laporan.', 'Hasil pengujian, screenshot, dokumentasi teknis, dan draft laporan.'],
], [3.0, 6.0, 5.0])
heading('3.4 Teknologi yang Dipelajari', 2)
bullet('Penerapan arsitektur MVC, migration, model, controller, route, middleware, dan policy pada Laravel.')
bullet('Pembuatan relasi database dan proses bisnis placement menggunakan service dan observer.')
bullet('Integrasi payment gateway, autentikasi eksternal, PDF, QR/kode verifikasi, dan ekspor Excel.')
bullet('Pembuatan interface responsif dengan Blade, Tailwind CSS, dan Alpine.js.')
heading('3.5 Kendala dan Solusi', 2)
add_table(['Kendala', 'Solusi'], [
    ['Penjadwalan dapat mengalami bentrok lokasi dan waktu.', 'Menambahkan validasi konflik jadwal pada backend dan policy akses.'],
    ['Penempatan peserta bergantung pada nilai dan kuota kelas.', 'Menggunakan service placement untuk mencocokkan rentang level dan kapasitas kursus.'],
    ['Integrasi layanan eksternal membutuhkan konfigurasi khusus.', 'Menguji konfigurasi secara bertahap dan memisahkan credential melalui environment.'],
    ['Dokumentasi fitur mudah menjadi terlalu panjang.', 'Meringkas database dengan ERD dan memusatkan detail teknis pada dokumentasi/lampiran.'],
], [6.0, 8.0])

# BAB IV
chapter('IV', 'ANALISIS DAN PERANCANGAN PROYEK')
heading('4.1 Identifikasi Masalah', 2)
para('Masalah utama yang menjadi dasar pembangunan ulang adalah kebutuhan terhadap sistem yang menghubungkan administrasi peserta, program, level, kursus, jadwal, pembelajaran, pembayaran, nilai, dan sertifikat. Sistem juga harus membedakan kewenangan tiap role dan menyediakan data yang dapat ditelusuri dari satu proses ke proses berikutnya.')
heading('4.2 Kebutuhan Sistem', 2)
add_table(['Jenis', 'Kebutuhan'], [
    ['Admin', 'Mengelola program, level, kursus, peserta, instruktur, lokasi, kelas, hari, jadwal, nilai, sertifikat, dan template.'],
    ['Instruktur', 'Melihat kursus terkait, mengisi risalah, absensi, dan nilai, serta melihat jadwal.'],
    ['Peserta', 'Melihat program, membuat pendaftaran, mengikuti placement, membayar, melihat kursus, riwayat, nilai/risalah, dan sertifikat.'],
    ['Sistem', 'Validasi role, konflik jadwal, kuota, alur placement, status pembayaran, dan verifikasi sertifikat.'],
], [4.0, 10.0])
heading('4.3 ERD Konseptual', 2)
para('ERD digunakan untuk melihat hubungan antar-entitas utama dalam sistem. Tabel database tidak ditulis satu per satu karena rincian field sudah tersedia pada migration dan dokumentasi teknis project. Diagram berikut merangkum entitas akun, profil pengguna, katalog kursus, pendaftaran, pembayaran, pembelajaran, penilaian, jadwal, dan sertifikat.')
add_table(['Entitas utama', 'Relasi utama', 'Fungsi'], [
    ['User - Peserta/Instruktur', 'User 1-1 profil', 'Akun login dan profil role.'],
    ['Program - Level - Kursus', 'Program 1-N Kursus; Level 1-N Kursus', 'Struktur katalog dan kelas kursus.'],
    ['Peserta - Pendaftaran', 'Peserta 1-N Pendaftaran', 'Pendaftaran program dan status proses.'],
    ['Pendaftaran - Score/Payment/Absensi', 'Pendaftaran 1-N', 'Placement, transaksi, dan kehadiran.'],
    ['Kursus - Jadwal/Risalah', 'Kursus 1-N', 'Pelaksanaan pertemuan dan dokumentasi pembelajaran.'],
    ['Template - Sertifikat', 'Template 1-N Sertifikat', 'Penerbitan dan verifikasi sertifikat.'],
], [5.0, 5.0, 4.0])
heading('4.3.1 Ringkasan Entitas ERD', 3)
para('Entitas utama yang ditampilkan pada diagram adalah users, pesertas, instrukturs, programs, levels, kursuses, pendaftarans, payments, jadwals, risalahs, absensis, scores, dan certificates. Entitas pendukung seperti lokasi, kelas, hari, template sertifikat, serta tabel penghubung digunakan untuk melengkapi proses bisnis dan ditunjukkan pada relasi/struktur project.', first_indent=False)
add_screenshot('erd-balai-kursus.png', 'Gambar 4.1 ERD konseptual aplikasi Balai Kursus')
heading('4.4 Perancangan Interface', 2)
para('Perancangan interface mengikuti pembagian role dan alur bisnis. Halaman publik dan autentikasi menjadi pintu masuk, dashboard menjadi pusat navigasi, sedangkan halaman CRUD memakai pola form, validasi, tabel, pencarian, dan pagination. Interface dirancang responsif menggunakan Tailwind CSS.')
add_table(['Kelompok halaman', 'Komponen rancangan'], [
    ['Publik dan autentikasi', 'Papan informasi, login, registrasi, lupa kata sandi, dan verifikasi sertifikat.'],
    ['Admin', 'Dashboard, sidebar menu, tabel master data, form, detail kursus, pembayaran, sertifikat.'],
    ['Instruktur', 'Dashboard kursus, jadwal read-only, risalah, absensi, dan nilai.'],
    ['Peserta', 'Katalog program, pendaftaran, pembayaran, kursus saya, riwayat, profil, dan sertifikat.'],
], [5.0, 9.0])
heading('4.4.1 Visualisasi Interface dari Project', 3)
para('Visual berikut diambil dari project yang dijalankan pada environment lokal. Gambar ini memperlihatkan bentuk halaman, warna, navigasi, dan komponen interface yang digunakan sebagai dasar pembahasan perancangan.', first_indent=False)
add_screenshot('tampilan-halaman-login.png', 'Gambar 4.2 Tampilan halaman login')
add_screenshot('tampilan-dashboard-admin.png', 'Gambar 4.3 Tampilan dashboard Admin')
add_screenshot('mockup-live-register.png', 'Gambar 4.4 Tampilan halaman registrasi')
add_screenshot('mockup-live-board.png', 'Gambar 4.5 Tampilan papan informasi publik')
add_screenshot('mockup-live-home.png', 'Gambar 4.6 Tampilan halaman login dari project berjalan')

# BAB V
chapter('V', 'IMPLEMENTASI DAN PENGUJIAN')
heading('5.1 Spesifikasi Teknologi', 2)
para('Project menggunakan PHP 8.1+, Laravel 10.10, MySQL, Laravel Modules, Blade, Vite, Tailwind CSS, Alpine.js, Midtrans, CAS/SSO, Dompdf, QR Code, dan Maatwebsite Excel. Environment lokal menggunakan Laragon. Konfigurasi credential layanan tidak ditulis dalam laporan.')
heading('5.2 Implementasi Fitur', 2)
para('Implementasi fitur dikelompokkan berdasarkan role dan proses utama agar pembahasan tetap ringkas. Screenshot berikut dipilih untuk mewakili fitur, sedangkan keseluruhan dokumentasi tampilan dicantumkan pada lampiran.')
add_screenshot('tampilan-manajemen-program.png', 'Gambar 5.1 Manajemen program')
add_screenshot('tampilan-manajemen-kursus.png', 'Gambar 5.2 Manajemen kursus')
add_screenshot('tampilan-manajemen-jadwal.png', 'Gambar 5.3 Manajemen jadwal')
add_screenshot('tampilan-pencatatan-absensi.png', 'Gambar 5.4 Pencatatan absensi')
add_screenshot('tampilan-penilaian-peserta.png', 'Gambar 5.5 Penilaian peserta')
add_screenshot('tampilan-proses-pembayaran-midtrans.png', 'Gambar 5.6 Proses pembayaran Midtrans')
add_screenshot('tampilan-manajemen-sertifikat.png', 'Gambar 5.7 Manajemen sertifikat')
add_screenshot('tampilan-verifikasi-sertifikat.png', 'Gambar 5.8 Verifikasi sertifikat')
heading('5.3 Pengujian', 2)
add_table(['No.', 'Skenario/Data Uji', 'Hasil yang Diharapkan', 'Hasil'], [
    ['1', 'Halaman login dibuka melalui browser lokal', 'Form login dan elemen navigasi tampil tanpa error', 'Berhasil; dibuktikan dengan screenshot UI lokal'],
    ['2', 'Halaman registrasi dibuka melalui browser lokal', 'Form registrasi dan validasi awal tampil', 'Berhasil; dibuktikan dengan screenshot UI lokal'],
    ['3', 'Papan informasi publik dibuka melalui browser lokal', 'Informasi publik dapat ditampilkan', 'Berhasil; dibuktikan dengan screenshot UI lokal'],
    ['4', 'Feature test autentikasi Laravel dijalankan', 'Test dapat membuat koneksi database pengujian', 'Belum dapat dijalankan; MySQL 127.0.0.1:3306 menolak koneksi'],
    ['5', 'Feature test pendaftaran, placement, dan pembayaran', 'Alur bisnis diuji dengan database bersih', 'Belum dapat diverifikasi karena database lokal belum aktif'],
    ['6', 'Feature test jadwal, sertifikat, dan otorisasi', 'Akses role dan data dibatasi sesuai aturan', 'Belum dapat diverifikasi karena database lokal belum aktif'],
    ['7', 'Screenshot fitur CRUD dan proses utama diperiksa', 'Tampilan fitur sesuai pembahasan BAB V', 'Berhasil; screenshot tersedia pada BAB V dan lampiran'],
    ['8', 'Pengujian online/deployment', 'Aplikasi dapat diakses melalui server publik', 'Belum dilakukan karena project belum dideploy'],
], [0.8, 5.0, 5.5, 2.7])
para('Pengujian yang dapat dilakukan pada environment saat penyusunan laporan adalah pemeriksaan tampilan melalui browser lokal dan pemeriksaan screenshot fitur. Feature test Laravel sudah tersedia pada folder tests, tetapi belum dapat dijalankan sampai selesai karena service MySQL pada 127.0.0.1:3306 tidak aktif. Pengujian database dan pengujian online harus diulang setelah database serta server tersedia.')
heading('5.4 Deployment', 2)
para('Deployment nyata belum diisi pada versi ini. Bagian berikut disiapkan untuk dilengkapi setelah informasi server tersedia:', first_indent=False)
add_table(['Item deployment', 'Isian yang diperlukan'], [
    ['URL aplikasi', '[Diisi]'], ['Platform/server', '[Diisi]'], ['Tanggal deployment', '[Diisi]'],
    ['Database produksi', '[Diisi]'], ['Status pengujian online', '[Diisi]'],
], [5.0, 9.0])
para('Kebutuhan teknis deployment meliputi Apache/Nginx, PHP 8.1+, MySQL, Composer, Node.js/NPM, konfigurasi .env, migrasi database, build asset Vite, storage, dan credential layanan eksternal.', first_indent=False)

# BAB VI
chapter('VI', 'PENUTUP')
heading('6.1 Kesimpulan', 2)
para('Aplikasi Balai Kursus berhasil dirancang dan dibangun sebagai platform manajemen kursus berbasis web dengan tiga role pengguna. Sistem mengintegrasikan katalog program, pendaftaran, placement dan penempatan kelas, pembayaran Midtrans, jadwal, risalah, absensi, penilaian, serta sertifikat yang dapat diverifikasi. Kegiatan PKL memberikan pengalaman dalam analisis kebutuhan, desain database dan interface, pengembangan Laravel, integrasi layanan, pengujian, dan dokumentasi.')
heading('6.2 Saran', 2)
bullet('Menambahkan notifikasi email atau notifikasi dalam aplikasi untuk jadwal, pembayaran, dan sertifikat.')
bullet('Mengembangkan laporan analitik operasional apabila kebutuhan bisnis sudah ditetapkan.')
bullet('Melakukan audit keamanan, backup, dan pemantauan performa secara berkala.')
bullet('Menyediakan deployment produksi dan pipeline CI/CD yang terdokumentasi.')

# Daftar pustaka
doc.add_page_break()
heading('DAFTAR PUSTAKA', 1)
para('Laravel. (2023). Laravel 10.x Documentation. https://laravel.com/docs/10.x. Diakses 24 Juli 2026.', first_indent=False)
para('Midtrans. (2024). Midtrans Documentation. https://docs.midtrans.com/. Diakses 24 Juli 2026.', first_indent=False)
para('Tailwind Labs. (2024). Tailwind CSS Documentation. https://tailwindcss.com/docs. Diakses 24 Juli 2026.', first_indent=False)
para('Vite Team. (2024). Vite Documentation. https://vite.dev/guide/. Diakses 24 Juli 2026.', first_indent=False)
para('PHP Group. (2024). PHP Manual. https://www.php.net/docs.php. Diakses 24 Juli 2026.', first_indent=False)
para('Nwidart. (2024). Laravel Modules Documentation. https://nwidart.com/laravel-modules. Diakses 24 Juli 2026.', first_indent=False)
para('Universitas Pendidikan Indonesia. (2024). Informasi Universitas Pendidikan Indonesia. https://www.upi.edu/. Diakses 24 Juli 2026.', first_indent=False)
para('Jay Lab Tech. (2026). Balai Kursus: Aplikasi Manajemen Kursus Berbasis Laravel. https://github.com/jay-lab-tech/balai-kursus. Diakses 24 Juli 2026.', first_indent=False)

# Lampiran
doc.add_page_break()
heading('LAMPIRAN', 1)
heading('Lampiran A. Daftar Screenshot', 2)
para('Screenshot berikut menjadi bukti visual dari project yang dijalankan pada environment lokal.', first_indent=False)
add_screenshot('mockup-live-home.png', 'Lampiran A.1 Halaman login')
add_screenshot('mockup-live-register.png', 'Lampiran A.2 Halaman registrasi')
add_screenshot('mockup-live-board.png', 'Lampiran A.3 Papan informasi publik')
heading('Lampiran B. Ringkasan Struktur Project', 2)
para('Struktur project terdiri dari app untuk model, controller, service, middleware, observer, policy, dan export; Modules untuk modul fitur; database untuk migration, seeder, dan factory; resources/views untuk view inti; routes untuk route web/API; public untuk asset; tests untuk pengujian; serta docs untuk dokumentasi.', first_indent=False)
heading('Lampiran C. Ringkasan Pengujian', 2)
para('Pengujian tampilan login, registrasi, dan papan informasi publik berhasil dilakukan pada browser lokal. Feature test yang membutuhkan database belum dapat dijalankan karena service MySQL lokal tidak aktif. Pengujian deployment belum dilakukan karena aplikasi belum dipasang pada server publik.', first_indent=False)
add_table(['Jenis pemeriksaan', 'Status', 'Catatan'], [
    ['Pemeriksaan UI lokal', 'Berhasil', 'Login, registrasi, dan papan informasi dapat ditangkap sebagai screenshot.'],
    ['Feature test Laravel', 'Tertahan', 'Koneksi MySQL 127.0.0.1:3306 ditolak saat test dijalankan.'],
    ['Pengujian online', 'Belum dilakukan', 'Menunggu deployment project.'],
], [5.0, 3.0, 6.0])
heading('Lampiran D. ERD dan Mockup', 2)
para('Diagram dan tampilan berikut merupakan visual yang digunakan dalam analisis dan perancangan project.', first_indent=False)
add_screenshot('erd-balai-kursus.png', 'Lampiran D.1 ERD konseptual')
add_screenshot('mockup-live-home.png', 'Lampiran D.2 Visualisasi interface')
heading('Lampiran E. Jurnal Kegiatan dan Pengesahan', 2)
para('Kegiatan PKL dilakukan dengan melanjutkan coding project Balai Kursus secara berkelompok dengan satu orang rekan. Pembagian pekerjaan mengikuti kesepakatan tim; penulis mengerjakan bagian coding lanjutan, integrasi fitur, pemeriksaan tampilan, dan dokumentasi, sedangkan rekan mengerjakan bagian fitur lain sesuai pembagian tugas dan melakukan koordinasi perubahan project.', first_indent=False)
add_table(['Periode', 'Kegiatan penulis', 'Koordinasi tim'], [
    ['Minggu 1–2', 'Membaca struktur project, melanjutkan coding, dan memahami alur data.', 'Menyepakati pembagian modul dan cara menggabungkan perubahan.'],
    ['Minggu 3–4', 'Mengembangkan dan memperbaiki fitur sesuai bagian yang dikerjakan.', 'Saling menguji hasil coding dan menyelesaikan konflik perubahan.'],
    ['Minggu 5–6', 'Melakukan pemeriksaan tampilan, pengujian yang tersedia, dan dokumentasi.', 'Meninjau hasil project dan menyepakati bahan laporan.'],
], [3.0, 6.0, 5.0])
para('Lembar pengesahan sekolah dan IDUKA tetap disediakan untuk diisi dengan pulpen, tanda tangan, tanggal, dan stempel secara manual.', first_indent=False)

for s in doc.sections:
    add_page_field(s)

doc.save(str(OUT))
print(f'Saved {OUT}')
