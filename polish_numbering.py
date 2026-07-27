from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path('docs/arsip-laporan/LAPORAN PKL RPL - VERSI VISUAL.docx')
OUT = Path('docs/LAPORAN PKL RPL - VERSI SIAP 2.docx')
doc = Document(str(SRC))

def naturalize(p, text):
    p.text = text
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Times New Roman')
        r._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 0)

def replace_start(prefix, text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            naturalize(p, text)
            return True
    return False

replacements = [
    ('Puji syukur penulis panjatkan', 'Puji syukur penulis panjatkan kepada Tuhan Yang Maha Esa karena laporan kegiatan PKL ini dapat diselesaikan. Laporan ini berisi kegiatan dan hasil pembangunan ulang aplikasi Balai Kursus di Direktorat Sistem dan Teknologi Informasi Universitas Pendidikan Indonesia.'),
    ('Penulis menyampaikan terima kasih', 'Penulis mengucapkan terima kasih kepada pihak sekolah, pembimbing, keluarga, dan semua pihak yang telah membantu selama kegiatan PKL. Penulis menyadari masih terdapat kekurangan dalam laporan ini, sehingga saran dan kritik sangat diharapkan.'),
    ('Perkembangan teknologi informasi mendorong', 'Pengelolaan kursus membutuhkan sistem yang dapat menyimpan data peserta, program, kelas, jadwal, pembayaran, dan hasil belajar secara teratur. Jika proses tersebut dilakukan secara terpisah, pencarian data dan penyusunan laporan menjadi lebih sulit.'),
    ('Aplikasi Balai Kursus dikembangkan', 'Aplikasi Balai Kursus dibuat untuk membantu proses tersebut. Di dalamnya terdapat akses untuk Admin, Instruktur, dan Peserta, mulai dari pendaftaran dan placement sampai pembayaran, kegiatan belajar, penilaian, dan sertifikat.'),
    ('Tujuan kegiatan ini adalah', 'Tujuan kegiatan ini adalah membangun aplikasi administrasi kursus berbasis web, menerapkan pengetahuan yang diperoleh di sekolah, dan memahami proses pengembangan perangkat lunak dalam lingkungan kerja.'),
    ('Masalah utama yang menjadi dasar pembangunan ulang', 'Masalah yang menjadi dasar pembangunan ulang adalah belum terhubungnya data peserta, program, level, kursus, jadwal, pembelajaran, pembayaran, nilai, dan sertifikat dalam satu alur kerja. Sistem juga harus membatasi akses sesuai role pengguna.'),
    ('ERD digunakan untuk menunjukkan', 'ERD digunakan untuk melihat hubungan antar-entitas utama dalam sistem. Laporan ini menampilkan relasi yang diperlukan untuk menjelaskan alur aplikasi, sedangkan rincian field tersedia pada dokumentasi teknis project.'),
    ('Perancangan interface mengikuti', 'Perancangan interface dibuat berdasarkan kebutuhan tiap role. Halaman menggunakan pola navigasi, form, tabel, validasi, dan detail data yang konsisten agar pengguna lebih mudah memahami alur aplikasi.'),
    ('Project menggunakan PHP', 'Project menggunakan PHP 8.1+, Laravel 10.10, MySQL, Laravel Modules, Blade, Vite, Tailwind CSS, Alpine.js, Midtrans, CAS/SSO, Dompdf, QR Code, dan Maatwebsite Excel. Laragon digunakan sebagai environment lokal. Credential layanan tidak dicantumkan dalam laporan.'),
    ('Implementasi fitur dikelompokkan', 'Fitur yang telah dibuat dijelaskan berdasarkan role dan proses utamanya. Gambar pada bagian ini merupakan contoh tampilan aplikasi; seluruh screenshot tambahan dicantumkan pada folder dokumentasi project.'),
    ('Project juga memiliki feature test', 'Pada source project tersedia pengujian fitur untuk autentikasi, otorisasi, pendaftaran, webhook pembayaran, sertifikat, dashboard, instruktur, dan konflik jadwal. Pengujian online tetap perlu dilakukan setelah aplikasi dipasang pada server tujuan.'),
    ('Aplikasi Balai Kursus berhasil', 'Berdasarkan hasil pengembangan, aplikasi Balai Kursus telah memiliki alur utama untuk mengelola program, pendaftaran, placement, kelas, pembayaran, jadwal, risalah, absensi, nilai, dan sertifikat. Pengembangan ini juga memberi pengalaman kepada penulis dalam menggunakan Laravel, membuat relasi database, menghubungkan layanan eksternal, dan melakukan pengujian.'),
    ('Screenshot fitur utama yang digunakan', 'Screenshot pada lampiran berisi tampilan fitur utama aplikasi, seperti autentikasi, dashboard, program, kursus, jadwal, risalah, absensi, nilai, pembayaran, sertifikat, dan profil. Screenshot dipilih sebagai bukti visual dari fitur yang dijelaskan pada BAB V.'),
    ('Struktur project terdiri dari', 'Struktur project terdiri dari folder app untuk model dan controller, Modules untuk fitur modular, database untuk migration dan seeder, resources/views untuk tampilan, routes untuk route web/API, tests untuk pengujian, serta docs untuk dokumentasi.'),
]
for prefix, text in replacements:
    replace_start(prefix, text)

# Make the retained BAB II material fit the numbering pattern of the new
# report. These headings are carried over from the original company profile,
# but they still belong to BAB II and should appear as sub-subsections.
heading_numbers = {
    'Periode Perguruan Tinggi Pendidikan Guru (PTPG)': '2.1.1 Periode Perguruan Tinggi Pendidikan Guru (PTPG)',
    'Periode FKIP UNPAD dan IKIP Bandung': '2.1.2 Periode FKIP UNPAD dan IKIP Bandung',
    'Periode Universitas Pendidikan Indonesia (UPI)': '2.1.3 Periode Universitas Pendidikan Indonesia (UPI)',
    'Visi': '2.2.1 Visi',
    'Misi': '2.2.2 Misi',
    'Tujuan': '2.2.3 Tujuan',
    '1. Bentuk Bulat (Lingkaran)': '2.3.1 Bentuk Bulat (Lingkaran)',
    '2. Pena Berbulu Lima dan Buku Terbuka': '2.3.2 Pena Berbulu Lima dan Buku Terbuka',
    '3. Sayap Burung yang Mengembang': '2.3.3 Sayap Burung yang Mengembang',
    '4. Kobaran Api di Pucuk Pena': '2.3.4 Kobaran Api di Pucuk Pena',
    '5. Perpaduan Warna Merah dan Putih': '2.3.5 Perpaduan Warna Merah dan Putih',
    'Visi DSTI UPI': '2.6.1 Visi DSTI UPI',
    'Misi DSTI UPI': '2.6.2 Misi DSTI UPI',
    'Tugas pokok dan fungsi': '2.6.3 Tugas Pokok dan Fungsi',
    '4.3.1 Daftar Entitas ERD Lengkap': '4.3.1 Ringkasan Entitas ERD',
    '4.4.1 Mockup Visual dari UI Project': '4.4.1 Visualisasi Interface dari Project',
}
for p in doc.paragraphs:
    key = p.text.strip()
    if key in heading_numbers:
        p.text = heading_numbers[key]
        p.paragraph_format.line_spacing = 1.5
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Times New Roman')
            r._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            r.font.size = Pt(12)
            r.font.bold = True

# Keep the static list of figures readable before Word updates fields.
replace_start('Gambar 4.1 ERD konseptual', 'Gambar 2.1 Sejarah Universitas Pendidikan Indonesia\nGambar 2.2 Logo Universitas Pendidikan Indonesia\nGambar 2.3 Direktorat Sistem dan Teknologi Informasi\nGambar 2.4 Struktur Organisasi DSTI\nGambar 4.1 ERD konseptual aplikasi\nGambar 4.2 Tampilan halaman login\nGambar 4.3 Tampilan dashboard Admin\nGambar 4.4 Tampilan halaman registrasi\nGambar 4.5 Tampilan papan informasi publik\nGambar 4.6 Tampilan halaman login dari project berjalan\nGambar 5.1-5.8 Screenshot implementasi fitur')

# Add section breaks: cover has no number; front matter uses lower Roman;
# BAB I starts a fresh Arabic sequence at page 1.
body = doc.element.body
body_sect = body.sectPr
template = deepcopy(body_sect)

def clear_footer_refs(sect):
    for child in list(sect):
        if child.tag == qn('w:footerReference'):
            sect.remove(child)

def set_page_number(sect, fmt, start):
    old = sect.find(qn('w:pgNumType'))
    if old is not None:
        sect.remove(old)
    pg = OxmlElement('w:pgNumType')
    pg.set(qn('w:fmt'), fmt)
    pg.set(qn('w:start'), str(start))
    sect.append(pg)

def add_break_after(paragraph, sect):
    ppr = paragraph._p.get_or_add_pPr()
    existing = ppr.find(qn('w:sectPr'))
    if existing is not None:
        ppr.remove(existing)
    ppr.append(deepcopy(sect))

cover_last = next(p for p in doc.paragraphs if p.text.strip() == '2026')
approval_first = next(p for p in doc.paragraphs if p.text.strip() == 'Lembar Pengesahan Sekolah')
bab1 = next(p for p in doc.paragraphs if p.text.strip() == 'BAB I')
before_bab1 = next(p for p in doc.paragraphs if p._p.getnext() is bab1._p)

cover_sect = deepcopy(template)
clear_footer_refs(cover_sect)
set_page_number(cover_sect, 'decimal', 1)
front_sect = deepcopy(template)
set_page_number(front_sect, 'lowerRoman', 1)
main_sect = deepcopy(template)
set_page_number(main_sect, 'decimal', 1)

add_break_after(cover_last, cover_sect)
add_break_after(before_bab1, front_sect)
# Reassert that the cover has no footer; the first numbered page is the
# approval section (Roman numeral i).
cover_break = cover_last._p.get_or_add_pPr().find(qn('w:sectPr'))
clear_footer_refs(cover_break)
body.remove(body_sect)
body.append(main_sect)

# Footer remains centered and contains a PAGE field. Cover has no footer ref;
# front matter and main content share the same footer part.
for section in doc.sections:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run('')
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.addnext(fld)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0, 0, 0)

# Request Word to refresh the table of contents when the document opens.
settings = doc.settings.element
update = settings.find(qn('w:updateFields'))
if update is None:
    update = OxmlElement('w:updateFields')
    settings.append(update)
update.set(qn('w:val'), 'true')

# Final OOXML cleanup: the first section (cover) must not reference the page
# number footer at all. The next section starts the Roman numbering.
first_section = body.xpath('./w:p/w:pPr/w:sectPr')[0]
for child in list(first_section):
    if child.tag.endswith('footerReference'):
        first_section.remove(child)

doc.save(str(OUT))
print(f'Saved {OUT}')
