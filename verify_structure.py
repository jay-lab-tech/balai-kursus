import sys
from docx import Document
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

p = Path('docs/LAPORAN KEGIATAN PKL RPL (UPI) - RAPI_FIXED.docx')
doc = Document(str(p))
print(f"Total paragraphs: {len(doc.paragraphs)}")
print('=' * 60)
for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    style = para.style.name if para.style else '?'
    if style.startswith('Heading') or style == 'heading' or style == 'Title' or txt.upper().startswith('BAB'):
        print(f'[{i}] ({style}) {txt}')
